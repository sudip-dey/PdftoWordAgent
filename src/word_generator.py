"""
Word document generation from extracted PDF content
"""

import logging
from pathlib import Path
from typing import List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.pdf_extractor import PDFPage, StyledText
from src.config import WordDocumentConfig

logger = logging.getLogger(__name__)


class WordGenerator:
    """Generate Word documents from extracted PDF content"""
    
    def __init__(self, config: WordDocumentConfig = None):
        """
        Initialize Word generator
        
        Args:
            config: WordDocumentConfig object for styling options
        """
        self.config = config or WordDocumentConfig()
    
    def generate(self, pages: List[PDFPage], output_path: str) -> str:
        """
        Generate Word document from PDF pages
        
        Args:
            pages: List of PDFPage objects containing styled text
            output_path: Path where Word document should be saved
            
        Returns:
            Path to generated Word document
            
        Raises:
            Exception: If document generation fails
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Generating Word document: {output_path}")
        
        try:
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(self.config.MARGIN_TOP)
                section.bottom_margin = Inches(self.config.MARGIN_BOTTOM)
                section.left_margin = Inches(self.config.MARGIN_LEFT)
                section.right_margin = Inches(self.config.MARGIN_RIGHT)
            
            # Process each page
            for page in pages:
                self._add_page_content(doc, page)
                
                # Add page break between pages (except after last page)
                if page.page_number < len(pages):
                    doc.add_page_break()
            
            # Save document
            doc.save(str(output_path))
            logger.info(f"Document saved successfully: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Failed to generate Word document: {e}")
            raise
    
    def _add_page_content(self, doc: Document, page: PDFPage) -> None:
        """
        Add content from a single PDF page to Word document
        
        Args:
            doc: python-docx Document object
            page: PDFPage object containing styled text
        """
        # Group text segments into paragraphs (split by newlines)
        paragraphs_data = []
        current_para = []
        
        for segment in page.text_segments:
            if "\n" in segment.text:
                # Split segment by newlines
                parts = segment.text.split("\n")
                for i, part in enumerate(parts):
                    if part:  # Skip empty strings
                        styled_part = StyledText(
                            text=part,
                            bold=segment.bold,
                            italic=segment.italic,
                            underline=segment.underline,
                            font_size=segment.font_size,
                            font_name=segment.font_name
                        )
                        current_para.append(styled_part)
                    
                    # Add paragraph after each newline except the last
                    if i < len(parts) - 1:
                        paragraphs_data.append(current_para)
                        current_para = []
            else:
                current_para.append(segment)
        
        # Add final paragraph
        if current_para:
            paragraphs_data.append(current_para)
        
        # Add paragraphs to document
        for para_segments in paragraphs_data:
            if not para_segments:
                # Empty paragraph for spacing
                doc.add_paragraph()
                continue
            
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = self.config.LINE_SPACING
            para.paragraph_format.space_before = Pt(self.config.SPACE_BEFORE)
            para.paragraph_format.space_after = Pt(self.config.SPACE_AFTER)
            
            # Add runs for each styled segment
            for segment in para_segments:
                run = para.add_run(segment.text)
                self._apply_styling(run, segment)
    
    def _apply_styling(self, run, segment: StyledText) -> None:
        """
        Apply styling to a text run
        
        Args:
            run: python-docx Run object
            segment: StyledText object with styling information
        """
        # Font properties
        run.font.name = self.config.DEFAULT_FONT_NAME
        run.font.size = Pt(self.config.DEFAULT_FONT_SIZE)
        run.font.color.rgb = RGBColor.from_string(self.config.DEFAULT_FONT_COLOR)
        
        # Apply text styling
        run.font.bold = segment.bold
        run.font.italic = segment.italic
        run.font.underline = segment.underline
        
        logger.debug(
            f"Applied styling - bold: {segment.bold}, "
            f"italic: {segment.italic}, underline: {segment.underline}"
        )
